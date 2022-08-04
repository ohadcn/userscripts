def excepthook(type, value, tb):
	import traceback
	traceback.print_exception(type, value, tb)
	input()
import sys
sys.excepthook = excepthook

try:
	from requests import get
	from docx import Document
except:
	from subprocess import check_call
	from sys import executable
	for package in ["requests", "python-docx"]:
		check_call([executable, "-m", "pip", "install", package])
	from requests import get
	from docx import Document

from csv import DictReader
from os import chdir, environ, stat, remove
# environ["PYTHONIOENCODING"] = "utf-8"
from os.path import isfile, join as pathJoin

from re import compile
from tempfile import gettempdir
from datetime import datetime
from time import sleep
temp_dir = gettempdir()

# chdir(temp_dir)

now = datetime.now().timestamp()
def get_csv(url):
	filename = pathJoin(temp_dir, url[url.rfind("/")+1:])
	if not isfile(filename) or now - stat(filename).st_mtime > 60 * 60 :
		print("downloading " + url)
		file_csv = get(url)
		# print(str(len(file_csv.text)) + " downloaded to " + filename)
		open(filename, "wt").write(file_csv.text)
	return DictReader(open(filename, "rt"))

def get_laws():
	return get_csv("https://production.oknesset.org/pipelines/data/bills/kns_bill/kns_bill.csv")

def get_docs():
	return get_csv("https://production.oknesset.org/pipelines/data/bills/kns_documentbill/kns_documentbill.csv")
	
def get_doc(url, retry = 3):
	filename = pathJoin(temp_dir, url[url.rfind("/")+1:])
	if(not isfile(filename)):
		print("downloading from " + url + " to " + filename)
		data = get(url)#, verify=False)
		#print(str(len(data.text)) + " downloaded from " + url + " to " + filename)
		#data.raw.decode_content = True
		open(filename, "wb").write(data.content)
	try:
		return Document(filename)
	except Exception as e:
		if retry > 0:
			print(filename + " failed to open, retrying")
			remove(filename)
			sleep(500)
			return get_doc(url, retry=retry-1)
		else:
			raise e
	
laws = {}
for law in get_laws():
	laws[law["BillID"]] = law

print(str(len(laws)) + " laws loaded")

dups=compile("(פ[/\\\\][0-9]+[/\\\\]2(1|3|4|2))")
init = compile("יוז(מות|מים|מת|ם):\t? +חבר(ות|י|ת)? הכנסת\t")
numbers = compile("[\d]+")
scored_laws = {}
for name in ["laws21", "laws22", "laws23"]:
	dict = DictReader(open(name+ ".csv", "rt"))
	for line in dict:
		if line.get("מספר חוק") and line.get("ניקוד לחוק") != None:
			scored_laws[line["מספר חוק"]] = line

scores = [['"שם הצעת החוק","מדרג","מספר חוק","ניקוד", "קישור להצעת החוק", "הסבר הדירוג","הערות אחרות","הגיע להצבעה?","עבר?","יוזם ראשון","חתומים"']] + [[]] * 5000
n = 1
CURRENT_KNESSET = "24"
for line in DictReader(open("laws" + CURRENT_KNESSET + ".csv", "rt")):
	if line.get("מספר חוק") and not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
		n+=1
		if not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
			n-=2
			if not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
				n+=1
				print(line.get("מספר חוק"), n, "no match")
			
	if not line.get("מספר חוק"):
		line["מספר חוק"] = ("פ/" + str(n) + "/" + CURRENT_KNESSET)
	if line.get("הסבר הדירוג").find("ראה חוק") > 0 and not line.get("מדרג") :
		line["מדרג"] = "dup_laws_bot"
	# print(line)
	scores[n] = [
		"\"" + line.get("שם הצעת החוק") + "\"", line.get("מדרג"), line.get("מספר חוק"),
		line.get("ניקוד לחוק") or line.get("ניקוד"),line.get("קישור להצעה"),"\""+line.get("הסבר הדירוג").replace("\"", "'")+"\"",
		line.get("הערות אחרות"),line.get("הגיע להצבעה?"),line.get("עבר?"),line.get("יוזם ראשון"),line.get("חתומים")]
	if line.get("ניקוד לחוק") != None:
		if line.get("מספר חוק"):
			scored_laws[line["מספר חוק"]] = line
	n+=1
# sys.exit(0)
duplicates = {}
laws_initiators = [[]]*5000

news_csv = 'קישור, שם, מספר\n'
unscored_csv = 'קישור, שם, מספר, עלה להצבעה\n'
old_csv = 'קישור, שם, מספר חדש, דירוג, מספר קודם\n'

split_initiators = compile("[\n\t]")
i = 0
laws_last = 0
for doc in get_docs():
	i += 1
	# if i%1000 == 0:
	# 	print(i)

	law = laws.get(doc["BillID"])
	if not law:
		print("law " + str(doc["BillID"]) + " not found")
		# print(doc)
		continue


	if law["KnessetNum"] != CURRENT_KNESSET:
		continue

	# 53 - הצעה ממשלתית
	# 54 - הצעה פרטית
	# 55 - ועדה
	if law["SubTypeID"] != '54':
		if law["SubTypeID"] not in ["53", "54", "55"]:
			print("unknown law type", law["SubTypeID"], law["SubTypeDesc"])
		continue
	
	# 1 - דיון מוקדם
	# 2 - הצעת חוק לקריאה הראשונה
	# 4 - הצעת חוק לקריאה השנייה והשלישית
	# 5 - הצעת חוק לקריאה השנייה והשלישית - לוח תיקונים
	# 8 - חוק - נוסח לא רשמי
	# 9 - חוק - פרסום ברשומות
	# 12 - מסמך מ.מ.מ
	# 17 - החלטת ממשלה
	# 46 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש
	# 51 - הצעת חוק לדיון מוקדם - נוסח מתוקן
	# 56 - הצעת חוק לקריאה הראשונה - נוסח לדיון בוועדה
	# 59 - חומר רקע
	# 101 - הצעת חוק לקריאה השניה והשלישית - פונצ  בננה
	# 102 - הצעת חוק לקריאה השניה והשלישית - לוח תיקונים - פונצ בננה
	# 103 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש- פונצ בננה
	if doc["GroupTypeID"] != '1':
		if doc["GroupTypeID"] not in ['51', "1", "2", "4", "5", "8", "9", "17", "46", "56", "59", "101", "102", "103", "12"]:
			print("unknown doc type", doc["GroupTypeID"], doc["GroupTypeDesc"])
		continue
	
	num = int(law["PrivateNumber"])
	laws_last += 1
	# if num != laws_last:
	# 	print(num, laws_last)
	law_name = "פ\\{}\\{}".format(CURRENT_KNESSET, num)
	if len(scores) <= num or not scores[num]:
		scores[num] = ["\"" + law["Name"].replace("\"", "'") + "\""] + [''] * 8
	if not scores[num][2]:
		scores[num][2] = law_name
	if not scores[num][4]:
		scores[num][4] = 'https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + law["BillID"]
	# print(law)
	if law["StatusID"] in ["104", "141"]:
		scores[num][7] = "0"
		scores[num][8] = "0"
	elif law["StatusID"] == "118":
		# חוק עבר
		scores[num][7] = "1"
		scores[num][8] = "1"
	elif law["StatusID"] == "177":
		if(law["PostponementReasonID"] in ["3013", "3012", "3011", "3010", "1065", "2511"]):
			# נדחה בגלל שהח"כ המציע עזב את הכנסת
			scores[num][7] = "0"
			scores[num][8] = "0"
		else:
			# נדחה בהצבעה
			scores[num][7] = "1"
			scores[num][8] = "0"
	else:
		scores[num][7] = "1"
		scores[num][8] = "0"

	old_names = None
	#print(doc["GroupTypeDesc"], doc["FilePath"])
	for p in get_doc(doc["FilePath"]).paragraphs:
		if init.match(p.text): #.find("יוזמים:      חברי הכנסת") == 0:
			initiators = [a.strip() for a in split_initiators.split(init.sub("", p.text).replace("_", ""))]
			initiators = [a for a in initiators if a]
			laws_initiators[num] = initiators
			scores[num][9:] = initiators

		other_names = [i[0] for i in dups.findall(p.text)]
		if p.text.find("חוק") > 0 and other_names and (p.text.find("זהות")>=0 or p.text.find("זהה")>=0):
			old_names = other_names
			for n in old_names:
				for v in [n, n.replace('/', '\\'), n.replace('\\', '/')]:
					if scored_laws.get(v):
						name = v
						if scored_laws[name]["שם הצעת החוק"].find(law['Name'][10:20]) == -1 and scored_laws[name]["שם הצעת החוק"].find(law['Name'][20:30]) == -1:
								print("כפילות בשם שונה!", law['Name'], law_name, scored_laws[name]["שם הצעת החוק"], name)
						if duplicates.get(law_name):
							#print("law scored already twice! " + str(scored_laws[name]) + str(duplicates[law_name]))
							if scored_laws[name]["ניקוד לחוק"] and duplicates[law_name]["ניקוד לחוק"] and scored_laws[name]["ניקוד לחוק"] != duplicates[law_name]["ניקוד לחוק"]:
								print("חוק דורג פעמיים בעבר בניקוד שונה", law["Name"], scored_laws[name]["ניקוד לחוק"], scored_laws[name]["מספר חוק"], duplicates[law_name]["ניקוד לחוק"], duplicates[law_name]["מספר חוק"])
						duplicates[law_name] = scored_laws[name]
						old_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + law["BillID"] + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "," + scored_laws[name]["ניקוד לחוק"] + "," + scored_laws[name]["מספר חוק"] + "\n")
						if not scores[int(law_name[5:])][3] and scored_laws[name]["ניקוד לחוק"] != "":
							scores[int(law_name[5:])][1:6] = ["dup_laws_bot", law_name, 
								scored_laws[name]["ניקוד לחוק"],'https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + law["BillID"],
								"\"" + (scored_laws[name]["הסבר הדירוג"] + " ראה חוק\n" + scored_laws[name]["מספר חוק"] + " " + scored_laws[name]["קישור להצעה"]).replace("\"", "'") + "\""]
						break

	if not old_names:
		news_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + law["BillID"] + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "\n")
		if scored_laws.get(name) == None and scores[num][3] == "":
			unscored_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + law["BillID"] + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "," + scores[num][7] + "\n")
	if not scores[num] or len(scores[num]) < 10 or not scores[num][9]:
		print("no iniitiators", [p.text for p in get_doc(doc["FilePath"]).paragraphs])

print(laws_last, len(scores), "laws")
open("unscored_laws.csv", "wt").write(unscored_csv)
open("new_laws.csv", "wt").write(news_csv)
open("scored_laws.csv", "wt").write(old_csv)
open("table.csv", "wt").write("\n".join([",".join([o or '' for o in line]) for line in scores if line]))
open("initiators.csv", "wt").write("\n".join([",".join([o or '' for o in line]) for line in laws_initiators]))

input()